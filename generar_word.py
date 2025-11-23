"""
Script para convertir el reporte markdown a formato Word
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Crear documento
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Leer archivo markdown
with open('REPORTE_METRICAS_COMPLETO.md', 'r', encoding='utf-8') as f:
    contenido = f.read()

# Función para agregar título
def agregar_titulo(texto, nivel=1):
    if nivel == 1:
        p = doc.add_heading(texto, level=1)
        p.runs[0].font.size = Pt(24)
        p.runs[0].font.color.rgb = RGBColor(0, 51, 153)
    elif nivel == 2:
        p = doc.add_heading(texto, level=2)
        p.runs[0].font.size = Pt(18)
        p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    elif nivel == 3:
        p = doc.add_heading(texto, level=3)
        p.runs[0].font.size = Pt(14)
        p.runs[0].font.color.rgb = RGBColor(51, 102, 153)
    else:
        p = doc.add_heading(texto, level=4)
        p.runs[0].font.size = Pt(12)

# Función para agregar párrafo con formato
def agregar_parrafo(texto):
    # Detectar texto en negrita
    if '**' in texto:
        p = doc.add_paragraph()
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if parte.startswith('**') and parte.endswith('**'):
                run = p.add_run(parte[2:-2])
                run.bold = True
            elif parte:
                p.add_run(parte)
    else:
        doc.add_paragraph(texto)

# Función para agregar lista
def agregar_lista(texto, numerada=False):
    if numerada:
        doc.add_paragraph(texto, style='List Number')
    else:
        doc.add_paragraph(texto, style='List Bullet')

# Procesar contenido línea por línea
lineas = contenido.split('\n')
en_bloque_codigo = False
codigo_buffer = []

for i, linea in enumerate(lineas):
    linea_limpia = linea.strip()
    
    # Bloques de código
    if linea_limpia.startswith('```'):
        if en_bloque_codigo:
            # Fin del bloque de código
            codigo_completo = '\n'.join(codigo_buffer)
            p = doc.add_paragraph(codigo_completo)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            codigo_buffer = []
            en_bloque_codigo = False
        else:
            # Inicio del bloque de código
            en_bloque_codigo = True
        continue
    
    if en_bloque_codigo:
        codigo_buffer.append(linea)
        continue
    
    # Títulos
    if linea_limpia.startswith('# '):
        agregar_titulo(linea_limpia[2:], 1)
    elif linea_limpia.startswith('## '):
        agregar_titulo(linea_limpia[3:], 2)
    elif linea_limpia.startswith('### '):
        agregar_titulo(linea_limpia[4:], 3)
    elif linea_limpia.startswith('#### '):
        agregar_titulo(linea_limpia[5:], 4)
    
    # Separadores
    elif linea_limpia == '---':
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)
        doc.add_paragraph()
    
    # Listas
    elif linea_limpia.startswith('- ') or linea_limpia.startswith('* '):
        agregar_lista(linea_limpia[2:], False)
    elif re.match(r'^\d+\.', linea_limpia):
        agregar_lista(re.sub(r'^\d+\.\s*', '', linea_limpia), True)
    
    # Checkboxes
    elif '- [ ]' in linea_limpia or '- [x]' in linea_limpia or '- [X]' in linea_limpia:
        texto = linea_limpia.replace('- [ ]', '☐').replace('- [x]', '☑').replace('- [X]', '☑')
        doc.add_paragraph(texto, style='List Bullet')
    
    # Líneas vacías
    elif not linea_limpia:
        continue
    
    # Párrafos normales
    else:
        agregar_parrafo(linea_limpia)

# Agregar encabezado de página
section = doc.sections[0]
header = section.header
header_para = header.paragraphs[0]
header_para.text = "DSS - Dashboard de Desempeño de Proyectos de Software"
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_para.runs[0].font.size = Pt(10)
header_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

# Agregar pie de página
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = "Reporte Completo de Métricas, KPIs y OLAP | Página "
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.runs[0].font.size = Pt(9)
footer_para.runs[0].font.color.rgb = RGBColor(100, 100, 100)

# Guardar documento
doc.save('REPORTE_METRICAS_COMPLETO.docx')
print("✅ Documento Word generado exitosamente: REPORTE_METRICAS_COMPLETO.docx")
